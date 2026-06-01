import io
import PyPDF2
from docx import Document
import re


def extract_text_from_pdf(file_bytes: io.BytesIO) -> str:
    """Extract text from PDF, handling multi-page documents."""
    try:
        reader = PyPDF2.PdfReader(file_bytes)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
        full_text = "\n\n".join(pages_text)
        return clean_text(full_text)
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {e}")


def extract_text_from_docx(file_bytes: io.BytesIO) -> str:
    """Extract text from DOCX including tables."""
    try:
        doc = Document(file_bytes)
        parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return clean_text("\n".join(parts))
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {e}")


def clean_text(text: str) -> str:
    """Clean extracted text for AI processing."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
    return text.strip()
